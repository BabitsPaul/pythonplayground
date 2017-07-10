import sys
from subprocess import CalledProcessError, check_output

from PyPDF2 import PdfFileReader
import docx
import os

# get parameters
if len(sys.argv) > 1:
    base = sys.argv[1:]
else:
    base = ['.']

# iterate over all base-folders
for b in base:
    print("Processing directory: ", b)
    os.chdir(b)

    # process all files in current base
    for fileName in os.listdir('.'):
        print("Processing file ", fileName)

        # get file extension
        tmp = os.path.splitext(fileName)
        if len(tmp) < 2:
            print("Error couldn't determine type of file ", fileName)
            continue
        ext = tmp[1]

        # create new file-name for pdf
        if ext == ".pdf":
            try:
                input1 = PdfFileReader(open(fileName, "rb"))
                nfn = input1.getDocumentInfo().title

                if nfn is None:
                    p = input1.getPage(1)
                    if p is None:
                        print("Error - Document doesn't contain first page")
                        continue

                    tmp = p.extractText()

                    if tmp is None:
                        print("Failed to extract text from document")
                        continue

                    nfn = tmp.split("\n")[0]
            except os.error:
                print("Failed to open file {}".format(b + "/" + fileName))
                continue
            except:
                print("Unknown error")
                continue
        #  get filename for docx (first paragraph of text)
        elif ext == ".docx":
            try:
                doc = docx.Document(fileName)

                if len(doc.paragraphs) == 0:
                    print("Error cant extract title from file {}".format(fileName))
                    continue

                nfn = doc.paragraphs[0].text
            except:
                print("Unknown error")
                continue
        # get filename for doc (requires catdoc)
        elif ext == ".doc":
            try:
                tmp = check_output(["catdoc", "-b", fileName], universal_newlines=True, stderr=sys.stdout)
                nfn = tmp.split("\n")[0]
            except CalledProcessError as e:
                print("Error - Failed to run catdoc for file {} error: {}".format(fileName, e))
                continue
        else:
            print("Cant process file-format: ", ext)
            continue

        if nfn is None:
            print("Error - failed to generate new file-name")
            continue

        # limit name to 255 chars (NFTS-max)
        if len(nfn) > 255 - len(ext):
            nfn = nfn[:256 - len(ext)]

        if len(nfn) == 0:
            print("No title found! Skipping file ", fileName)
            continue

        # append file-extension
        nfn += ext

        # rename file
        try:
            os.rename(fileName, nfn)
            print("File renamed from {} to {}".format(fileName, nfn))
        except os.error:
            print("Failed to rename file {}".format(fileName))
